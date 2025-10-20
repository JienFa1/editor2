# -*- coding: utf-8 -*-
import sys

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

"""
Utilities for reading DOCX files referenced by Config.DOCUMENT.
Preserves paragraph ordering so we can later map edited text back.
"""

from dataclasses import dataclass
from typing import List, Optional, Sequence, Tuple
import os

from docx import Document

from . import Config


@dataclass
class ParagraphRecord:
    """Paragraph content paired with its original index inside the DOCX."""

    docx_index: int
    text: str


def get_document_path() -> str:
    """Return the DOCX path from Config.DOCUMENT after validating it."""
    path = getattr(Config, "DOCUMENT", None)
    if not isinstance(path, str) or not path.strip():
        raise RuntimeError("Config.DOCUMENT must be a non-empty string path to a .docx file.")
    path = path.strip()
    if not os.path.isfile(path):
        raise RuntimeError(f"DOCX file not found: {path}")
    return path


def read_paragraphs_from_config() -> List[str]:
    """Return raw paragraph texts from the configured DOCX."""
    path = get_document_path()
    doc = Document(path)
    return [(p.text or "").strip() for p in doc.paragraphs]


def paragraphs_to_big_text(paragraphs: List[str]) -> str:
    """Join paragraphs into a single string separated by double newlines."""
    cleaned = [p.strip() for p in paragraphs if isinstance(p, str)]
    return "\n\n".join(cleaned)


def document_to_big_text() -> str:
    """Read the configured DOCX and return a concatenated text string."""
    return paragraphs_to_big_text(read_paragraphs_from_config())


def extract_textual_paragraphs(doc: Document, *, keep_empty: bool = False) -> List[ParagraphRecord]:
    """
    Return a list of ParagraphRecord, preserving the original paragraph order.

    Args:
        doc: The python-docx Document.
        keep_empty: When True, include empty paragraphs as placeholders.
    """
    records: List[ParagraphRecord] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").strip()
        if not text and not keep_empty:
            continue
        records.append(ParagraphRecord(docx_index=idx, text=text))
    return records


def load_document_with_text(path: Optional[str] = None) -> Tuple[Document, List[ParagraphRecord]]:
    """
    Load a DOCX file and return the Document plus textual paragraphs with indices.

    Args:
        path: Optional override path. Defaults to Config.DOCUMENT.
    """
    target_path = path or get_document_path()
    if not os.path.isfile(target_path):
        raise RuntimeError(f"DOCX file not found: {target_path}")
    doc = Document(target_path)
    paragraphs = extract_textual_paragraphs(doc)
    return doc, paragraphs


def document_to_big_text_with_mapping() -> Tuple[str, Document, List[ParagraphRecord]]:
    """
    Read Config.DOCUMENT and return big_text together with Document and paragraph map.

    Returns:
        big_text: Concatenated textual paragraphs.
        document: Loaded python-docx Document object.
        paragraphs: ParagraphRecord list preserving DOCX ordering.
    """
    document, paragraphs = load_document_with_text()
    lines: Sequence[str] = [item.text for item in paragraphs]
    big_text = paragraphs_to_big_text(list(lines))
    return big_text, document, paragraphs
